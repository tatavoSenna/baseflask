import pytest
from docx import Document as DocxDocument
from io import BytesIO
from unittest.mock import patch, Mock
from app.test import factories
from app.test import factories
from flask import jsonify
from datetime import datetime, timedelta
from app.models.documents import Document, DocumentTemplate
from app.documents.controllers import (
    create_document_controller,
    document_creation_email_controller,
)
from app.serializers.document_serializers import DocumentListSerializer


def generate_docx(text: str = "") -> BytesIO:
    docx_document = DocxDocument()
    docx_document.add_paragraph(text)
    docx = BytesIO()
    docx_document.save(docx)
    return docx


@pytest.fixture
def user():
    return factories.UserFactory()


@pytest.fixture
def document_template():
    def _document_template(text_type, fields):

        if not isinstance(fields, list):
            fields = [fields]
        variables = {}

        # builds the variables field of the document template model
        for field in fields:
            if field.get("variable", None):
                variables[field["variable"]["name"]] = field["variable"]

        return factories.DocumentTemplateFactory(
            text_type=text_type,
            form=[{"fields": fields}],
            workflow={"nodes": {}, "created_by": "", "current_node": "0"},
            variables=variables,
        )

    return _document_template


@patch("app.documents.controllers.RemoteDocument.download_docx_from_template")
@patch("app.documents.controllers.RemoteDocument.upload_filled_docx_to_documents")
@patch("app.documents.controllers.convert_docx_to_pdf_and_save")
def test_creation_of_docx_document_with_date_field(
    convert_docx_to_pdf_and_save,
    upload_filled_docx_to_documents,
    download_docx_from_template,
    user,
    document_template,
):
    date_field = {
        "info": "",
        "type": "date",
        "label": "",
        "variable": {"name": "DATA", "type": "date", "doc_display_style": "%d/%m/%Y"},
    }

    template_docx = generate_docx("Data: {{ DATA }}")
    download_docx_from_template.return_value = template_docx

    template = document_template(".docx", date_field)
    variables = {"DATA": "2021-12-31T12:34:56.789Z"}

    document = create_document_controller(
        user_id=user.id,
        user_email=user.email,
        company_id=user.company.id,
        document_template_id=template.id,
        title="title",
        username=user.username,
        received_variables=variables,
        parent_id=None,
        is_folder=False,
    )

    upload_filled_docx_to_documents.assert_called_once()
    upload_filled_docx_to_documents_args = (
        upload_filled_docx_to_documents.call_args.args
    )

    assert upload_filled_docx_to_documents_args[0] == document
    assert (
        DocxDocument(upload_filled_docx_to_documents_args[1]).paragraphs[0].text
        == "Data: 31/12/2021"
    )
    assert upload_filled_docx_to_documents_args[2] == ".docx"


@patch("app.documents.controllers.RemoteDocument.download_docx_from_template")
@patch("app.documents.controllers.RemoteDocument.upload_filled_docx_to_documents")
@patch("app.documents.controllers.convert_docx_to_pdf_and_save")
def test_creation_of_docx_document_with_number_field_with_doc_display_style(
    convert_docx_to_pdf_and_save,
    upload_filled_docx_to_documents,
    download_docx_from_template,
    user,
    document_template,
):
    """Tests if the NumberFormater respects the doc_display_string when outputing the variable
    The 'doc_display_style' property of hte variable has been deprecated for number, but we
    need this test for back compatibillity of templates.

    Args:
        convert_docx_to_pdf_and_save (MagicMock): to avoid calling convert_api
        upload_filled_docx_to_documents (MagicMock: to avoid calling aws s3
        download_docx_from_template (MagicMock): to avoid calling aws s3
        user (models.user.User): Fixture to represent the user that creates the document
        document_template (Function): Function to generate a Template model
    """

    number_field = {
        "info": "",
        "type": "number",
        "label": "",
        "variable": {
            "name": "NUMBER",
            "type": "number",
            "doc_display_style": "extended",
        },
    }

    template_docx = generate_docx("{{ NUMBER }}")
    download_docx_from_template.return_value = template_docx

    template = document_template(".docx", number_field)
    variables = {"NUMBER": "11"}

    document = create_document_controller(
        user_id=user.id,
        user_email=user.email,
        company_id=user.company.id,
        document_template_id=template.id,
        title="title",
        username=user.username,
        received_variables=variables,
        parent_id=None,
        is_folder=False,
    )

    upload_filled_docx_to_documents.assert_called_once()
    upload_filled_docx_to_documents_args = (
        upload_filled_docx_to_documents.call_args.args
    )

    assert upload_filled_docx_to_documents_args[0] == document
    assert (
        DocxDocument(upload_filled_docx_to_documents_args[1]).paragraphs[0].text
        == "11 (onze)"
    )
    assert upload_filled_docx_to_documents_args[2] == ".docx"


@patch("app.documents.controllers.RemoteDocument.download_docx_from_template")
@patch("app.documents.controllers.RemoteDocument.upload_filled_docx_to_documents")
@patch("app.documents.controllers.convert_docx_to_pdf_and_save")
@patch("app.documents.formatters.variables_formatter.requests.get")
def test_creation_of_docx_document_with_database_field(
    requests_get,
    convert_docx_to_pdf_and_save,
    upload_filled_docx_to_documents,
    download_docx_from_template,
    user,
    document_template,
):
    mock_response = [
        {"userId": 1, "id": 1, "title": "Mock Response", "completed": False}
    ]

    requests_get.return_value = Mock(ok=True)
    requests_get.return_value.json.return_value = mock_response

    database_field = {
        "info": "",
        "type": "database",
        "label": "",
        "variable": {
            "name": "KEY_VALUE",
            "type": "database",
            "search_key": "id",
            "database_endpoint": "http://jsonplaceholder.typicode.com/todos",
        },
    }

    template_docx = generate_docx("{{ KEY_VALUE.TITLE }} {{ KEY_VALUE.COMPLETED }}")
    download_docx_from_template.return_value = template_docx

    template = document_template(".docx", database_field)
    variables = {"KEY_VALUE": "1"}

    document = create_document_controller(
        user_id=user.id,
        user_email=user.email,
        company_id=user.company.id,
        document_template_id=template.id,
        title="title",
        username=user.username,
        received_variables=variables,
        parent_id=None,
        is_folder=False,
    )

    upload_filled_docx_to_documents.assert_called_once()
    upload_filled_docx_to_documents_args = (
        upload_filled_docx_to_documents.call_args.args
    )

    assert upload_filled_docx_to_documents_args[0] == document
    assert (
        DocxDocument(upload_filled_docx_to_documents_args[1]).paragraphs[0].text
        == "Mock Response False"
    )
    assert upload_filled_docx_to_documents_args[2] == ".docx"


@patch("app.documents.controllers.RemoteDocument.download_docx_from_template")
@patch("app.documents.controllers.RemoteDocument.upload_filled_docx_to_documents")
@patch("app.documents.controllers.convert_docx_to_pdf_and_save")
def test_creation_of_docx_document_with_number_field_with_out_doc_display_style(
    convert_docx_to_pdf_and_save,
    upload_filled_docx_to_documents,
    download_docx_from_template,
    user,
    document_template,
):
    """Test the many different formats generated by the NumberFormatter

    Args:
        convert_docx_to_pdf_and_save (MagicMock): to avoid calling convert_api
        upload_filled_docx_to_documents (MagicMock: to avoid calling aws s3
        download_docx_from_template (MagicMock): to avoid calling aws s3
        user (models.user.User): Fixture to represent the user that creates the document
        document_template (Function): FIxture, a function to generate a Template model
    """

    number_field = {
        "info": "",
        "type": "number",
        "label": "",
        "variable": {
            "name": "NUMBER",
            "type": "number",
        },
    }

    template_docx = generate_docx(
        "{{ NUMBER }} {{ NUMBER.extended }} {{ NUMBER.ordinal}} {{ NUMBER.percent }}"
    )
    download_docx_from_template.return_value = template_docx

    template = document_template(".docx", number_field)
    variables = {"NUMBER": "11"}

    document = create_document_controller(
        user_id=user.id,
        user_email=user.email,
        company_id=user.company.id,
        document_template_id=template.id,
        title="title",
        username=user.username,
        received_variables=variables,
        parent_id=None,
        is_folder=False,
    )

    upload_filled_docx_to_documents.assert_called_once()
    upload_filled_docx_to_documents_args = (
        upload_filled_docx_to_documents.call_args.args
    )

    assert upload_filled_docx_to_documents_args[0] == document
    assert (
        DocxDocument(upload_filled_docx_to_documents_args[1]).paragraphs[0].text
        == "11 11 (onze) décimo primeiro 11% (onze por cento)"
    )
    assert upload_filled_docx_to_documents_args[2] == ".docx"


class TestDocumentCreation:
    @pytest.fixture
    def test_company(self):
        return factories.CompanyFactory(id=17)

    @pytest.fixture
    def test_user(self, test_company):
        return factories.UserFactory(company=test_company, email="testemail@gmail.com")

    @pytest.fixture
    def test_workflow(self):
        workflow = {
            "nodes": {
                "544": {
                    "next_node": "5521",
                    "changed_by": "",
                    "title": "step1",
                    "deadline": 1,
                },
                "5521": {"next_node": "3485", "changed_by": "", "title": "step2"},
            },
            "current_node": "544",
            "created_by": "",
        }
        return workflow

    @pytest.fixture
    def test_empty_workflow(self):
        empty_workflow = {"nodes": {}, "created_by": "", "current_node": "0"}
        return empty_workflow

    @pytest.fixture
    def test_form(self):
        form = [
            {
                "title": "teste dummy",
                "fields": [
                    {
                        "info": "",
                        "type": "text",
                        "label": "Dummy text",
                        "variable": {
                            "name": "DUMMY",
                            "type": "string",
                            "doc_display_style": "plain",
                        },
                    }
                ],
            }
        ]
        return form

    @pytest.fixture
    def test_variables(self):
        variables = {"DUMMY": "test text"}
        return variables

    @pytest.fixture
    def test_txt_template(self, test_company, test_workflow, test_form):
        return factories.DocumentTemplateFactory(
            company=test_company,
            workflow=test_workflow,
            text_type=".txt",
            form=test_form,
        )

    @pytest.fixture
    def test_txt_template_without_workflow(
        self, test_company, test_empty_workflow, test_form
    ):
        return factories.DocumentTemplateFactory(
            company=test_company,
            workflow=test_empty_workflow,
            text_type=".txt",
            form=test_form,
        )

    @patch("app.documents.controllers.RemoteDocument.upload_filled_text_to_documents")
    @patch("app.documents.controllers.RemoteDocument.download_text_from_template")
    @patch("app.documents.controllers.fill_text_with_variables")
    @patch("app.documents.controllers.document_creation_email_controller")
    def test_send_email_on_document_creation_with_workflow(
        self,
        document_creation_email_controller,
        fill_text_with_variables,
        download_text_from_template,
        upload_filled_text_to_documents,
        test_user,
        test_company,
        test_txt_template,
        test_variables,
    ):

        document_title = "default title"

        document_txt = create_document_controller(
            test_user.id,
            test_user.email,
            test_user.company.id,
            test_txt_template.id,
            document_title,
            test_user.name,
            test_variables,
            None,
            False,
        )

        document_creation_email_controller.assert_called_once_with(
            document_title, test_company.id
        )

    @patch("app.documents.controllers.RemoteDocument.upload_filled_text_to_documents")
    @patch("app.documents.controllers.RemoteDocument.download_text_from_template")
    @patch("app.documents.controllers.fill_text_with_variables")
    @patch("app.documents.controllers.document_creation_email_controller")
    def test_send_email_on_document_creation_without_workflow(
        self,
        document_creation_email_controller,
        fill_text_with_variables,
        download_text_from_template,
        upload_filled_text_to_documents,
        test_user,
        test_txt_template_without_workflow,
        test_variables,
    ):

        document_title = "default title"

        document_txt = create_document_controller(
            test_user.id,
            test_user.email,
            test_user.company.id,
            test_txt_template_without_workflow.id,
            document_title,
            test_user.name,
            test_variables,
            None,
            False,
        )

        document_creation_email_controller.assert_not_called()

    @patch("app.documents.controllers.RemoteDocument.upload_filled_text_to_documents")
    @patch("app.documents.controllers.RemoteDocument.upload_filled_docx_to_documents")
    @patch("app.documents.controllers.RemoteDocument.download_text_from_template")
    @patch("app.documents.controllers.RemoteDocument.download_docx_from_template")
    @patch("app.documents.controllers.convert_docx_to_pdf_and_save")
    @patch("app.documents.controllers.fill_text_with_variables")
    @patch("app.documents.controllers.fill_docx_with_variables")
    @patch("app.documents.controllers.document_creation_email_controller")
    def test_create_document(
        self,
        mock_email_controller,
        upload_filled_text_to_documents,
        upload_filled_docx_to_documents,
        download_text_from_template,
        download_docx_from_template,
        convert_docx_to_pdf_and_save,
        fill_text_with_variables,
        fill_docx_with_variables,
        test_workflow,
        test_form,
        test_variables,
    ):

        company = factories.CompanyFactory(id=17)
        user = factories.UserFactory(company=company, email="testemail@gmail.com")

        document_template_txt = factories.DocumentTemplateFactory(
            company=company, workflow=test_workflow, text_type=".txt", form=test_form
        )

        document_template_docx = factories.DocumentTemplateFactory(
            company=company, workflow=test_workflow, text_type=".docx", form=test_form
        )

        document_title = "default title"

        variables = {}

        assert (
            DocumentTemplate.query.get(document_template_txt.id).company_id
            == company.id
        )

        assert (
            DocumentTemplate.query.get(document_template_docx.id).company_id
            == company.id
        )

        document_txt = create_document_controller(
            user.id,
            user.email,
            user.company_id,
            document_template_txt.id,
            document_title,
            user.name,
            variables,
            None,
            False,
        )

        document_docx = create_document_controller(
            user.id,
            user.email,
            user.company_id,
            document_template_docx.id,
            document_title,
            user.name,
            variables,
            None,
            False,
        )

        assert document_txt.user_id == user.id
        assert document_txt.company_id == company.id
        assert document_txt.document_template_id == document_template_txt.id
        assert document_txt.title == document_title
        assert document_txt.workflow["created_by"] == user.name
        assert document_txt.text_type == ".txt"
        assert (
            datetime.strptime(
                document_txt.workflow["nodes"]["544"]["due_date"],
                "%Y-%m-%d %H:%M:%S.%f",
            ).day
            == (datetime.today() + timedelta(days=1)).day
        )
        assert document_txt.due_date.day == (datetime.today() + timedelta(days=1)).day

        assert document_docx.document_template_id == document_template_docx.id
        assert document_docx.text_type == ".docx"
        assert (
            datetime.strptime(
                document_docx.workflow["nodes"]["544"]["due_date"],
                "%Y-%m-%d %H:%M:%S.%f",
            ).day
            == (datetime.today() + timedelta(days=1)).day
        )
        assert document_docx.due_date.day == (datetime.today() + timedelta(days=1)).day

    @patch("app.documents.controllers.send_email_controller")
    def test_email_create_document(self, send_email_on_document_creation_mock):
        company_id = 134
        company = factories.CompanyFactory(id=company_id)
        user = factories.UserFactory(company=company, email="testemail@gmail.com")
        user2 = factories.UserFactory(company=company, email="testemail2@gmail.com")
        email_title = "teste"

        email_list = []
        email_list.append(user.email)
        email_list.append(user2.email)

        document_creation_email_controller(email_title, company_id)

        send_email_on_document_creation_mock.assert_called_once_with(
            "leon@lawing.com.br",
            email_list,
            "New Document created",
            email_title,
            "d-83efa7b8d2fb4742a69dd9059324e148",
        )
