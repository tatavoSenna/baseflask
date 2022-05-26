import json
import logging
import base64
import io
import boto3
import requests
import copy
import convertapi

from app import jinja_env
from bs4 import BeautifulSoup
from datetime import datetime, timedelta, date
from docxtpl import DocxTemplate
from docx import Document as docxDocument
from docx.shared import Cm
from flask import current_app, jsonify
from operator import ne
from PIL import Image
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import Mail
from sqlalchemy import desc, nulls_last, select
from unittest.mock import MagicMock
from werkzeug.exceptions import Unauthorized, BadRequest

from app import db
from app.constants import months
from app.models.documents import Document, DocumentTemplate
from app.models.company import Company, Webhook
from app.models.user import User, Group
from app.workflow.services import (
    format_workflow_responsible_group,
    format_workflow_responsible_users,
)
from app.documents.formatters.variables_formatter import (
    format_variables,
)
from .remote import RemoteDocument


def get_document_template_list_controller(company_id):

    document_templates = (
        DocumentTemplate.query.filter_by(company_id=company_id, deleted=False)
        .filter_by(published=True)
        .order_by(
            nulls_last(desc(DocumentTemplate.favorite)),
            desc(DocumentTemplate.created_at),
        )
        .all()
    )
    return document_templates


def get_document_template_details_controller(company_id, template_id):

    document_template = DocumentTemplate.query.filter_by(
        company_id=company_id, id=template_id
    ).first()

    return document_template


def create_document_controller(
    user_id,
    user_email,
    company_id,
    document_template_id,
    title,
    username,
    received_variables,
    parent_id,
    is_folder,
    draft=False,
):
    # Get document template from the database using template_id and company_id
    document_template = DocumentTemplate.query.get(document_template_id)
    document_company = Company.query.filter_by(id=company_id).first()

    if document_company.remaining_documents <= 0:
        raise Unauthorized(
            description="You have reached the document limit offered by your plan"
        )

    # Get global variables
    nome_contrato = received_variables.get("NOME_CONTRATO", None)
    data_inicio_contrato = received_variables.get("DATA_INICIO_CONTRATO", None)
    data_final_contrato = received_variables.get("DATA_FINAL_CONTRATO", None)
    data_assinatura = received_variables.get("DATA_ASSINATURA", None)
    valor_contrato = received_variables.get("VALOR_CONTRATO", None)

    # Specify version 0 to create Document object
    current_date = datetime.now().astimezone().replace(microsecond=0).isoformat()
    current_user = User.query.filter_by(id=user_id).first()
    if draft:
        version = [
            {
                "description": "Rascunho",
                "email": user_email,
                "created_at": current_date,
                "id": "-1",
                "comments": None,
                "created_by": current_user.name,
            }
        ]
        title = "Rascunho: " + title
    else:
        version = [
            {
                "description": "Version 0",
                "email": user_email,
                "created_at": current_date,
                "id": "0",
                "comments": None,
                "created_by": current_user.name,
            }
        ]
    # Get workflow from the template and update it
    document_workflow = document_template.workflow
    document_workflow["created_by"] = username
    current_node = document_workflow["current_node"]
    try:
        document_workflow["nodes"][current_node]["due_date"] = str(
            datetime.today()
            + timedelta(days=document_workflow["nodes"][current_node]["deadline"])
        )
        due_date = document_workflow["nodes"][current_node]["due_date"]
    except:
        due_date = None
    try:
        step_name = document_workflow["nodes"][current_node]["title"]
    except KeyError:
        step_name = None

    # Create Document object
    document = Document(
        user_id=user_id,
        company_id=company_id,
        form=document_template.form,
        workflow=document_workflow,
        signers=document_template.signers,
        variables=received_variables,
        versions=version,
        created_at=datetime.now().astimezone().replace(microsecond=0).isoformat(),
        title=title,
        current_step=step_name,
        document_template_id=document_template_id,
        text_type=document_template.text_type,
        nome_contrato=nome_contrato,
        data_inicio_contrato=data_inicio_contrato,
        data_final_contrato=data_final_contrato,
        data_assinatura=data_assinatura,
        valor_contrato=valor_contrato,
        parent_id=parent_id,
        due_date=due_date,
        draft=draft,
    )
    # Add document to the database
    db.session.add(document)
    db.session.commit()

    if not draft:
        try:
            create_remote_document(
                document, received_variables, document_template, company_id
            )
        except:
            delete_document_controller(document)
            raise

    """ Disable the email while we don't have the correct sendgrid key 
    Will migrate this to a lambda with the new document sns hook
    # Send email informing document creation
    if step_name != None:
        try:
            document_creation_email_controller(title, company_id)
        except Exception as e:
            logging.exception(
                "Failed to send emails on document creation. One or more emails is bad formated or invalid"
            )
    """
    return document


def create_folder_controller(
    user_id, user_email, company_id, title, username, parent_id, is_folder
):

    document = Document(
        user_id=user_id,
        company_id=company_id,
        created_at=datetime.now().astimezone().replace(microsecond=0).isoformat(),
        title=title,
        parent_id=parent_id,
        is_folder=is_folder,
    )
    db.session.add(document)
    db.session.commit()

    return document


def get_current_date_dict():
    date = {}

    now = datetime.now()
    date["day"] = str(now.day).zfill(2)
    date["month"] = months[now.month - 1]
    date["year"] = now.year
    date["today"] = f'{ date["day"] }/{ date["month"] }/{ date["year"] }'

    return date


def get_document_controller(document_id):
    document = Document.query.filter_by(id=document_id).first()
    return document


def delete_document_controller(document):
    if not isinstance(document, MagicMock):
        document.deleted = True
        document.deleted_at = datetime.utcnow()
        db.session.commit()


def get_document_version_controller(document_id):
    document = get_document_controller(document_id)
    # take size of array with will be version + 1
    current_version = document.versions[0]["id"]
    return current_version


def get_comments_version_controller(document_id):
    document = get_document_controller(document_id)
    current_comments = document.versions[0]["comments"]
    return current_comments


def save_signers_controller(document_id, signers_variables):
    document = get_document_controller(document_id)
    # need to make a copy so that the 'signers' and 'variables' JSON changes are tracked
    variables = copy.deepcopy(document.variables)
    variables.update(signers_variables)
    document.variables = variables
    db.session.add(document)
    db.session.commit()

    return document


def create_new_version_controller(document_id, description, user_email, comments):
    document = Document.query.filter_by(id=document_id).first()
    current_user = User.query.filter_by(email=user_email).first()
    versions = document.versions
    current_version = int(versions[0]["id"])
    new_version = current_version + 1
    current_date = datetime.now().astimezone().replace(microsecond=0).isoformat()
    version = {
        "description": description,
        "email": user_email,
        "created_at": current_date,
        "id": str(new_version),
        "comments": comments,
        "created_by": current_user.name,
    }

    # need to make a copy to track changes to JSON, otherwise the changes are not updated
    document.versions = copy.deepcopy(document.versions)
    document.versions.insert(0, version)
    db.session.add(document)
    db.session.commit()


def download_document_text_controller(document_id, version_id):
    document = get_document_controller(document_id)
    remote_document = RemoteDocument()
    textfile = remote_document.download_text_from_documents(
        document, version_id
    ).decode("utf-8")
    return textfile


def download_document_docx_controller(document_id, version_id):
    document = get_document_controller(document_id)
    remote_document = RemoteDocument()
    docx_io = remote_document.download_docx_from_documents(document, version_id)
    return docx_io


def upload_document_text_controller(document_id, document_text):
    document = get_document_controller(document_id)
    # convert from string to bytes object
    document_text = bytearray(document_text, encoding="utf8")
    remote_document = RemoteDocument()
    remote_document.upload_filled_text_to_documents(document, document_text)


def next_status_controller(document_id, username):
    document = get_document_controller(document_id)
    workflow = document.workflow

    current_node = workflow["current_node"]
    if current_node is None:
        return document, 1

    next_node = workflow["nodes"][current_node]["next_node"]
    # check if there is a next node
    if next_node is None:
        return document, 1

    # need to make a copy to track changes to JSON, otherwise the changes are not updated
    document.workflow = copy.deepcopy(document.workflow)
    document.workflow["nodes"][workflow["current_node"]]["changed_by"] = username
    document.workflow["current_node"] = next_node
    document.current_step = document.workflow["nodes"][next_node]["title"]
    try:
        document.workflow["nodes"][next_node]["due_date"] = str(
            datetime.today()
            + timedelta(days=document.workflow["nodes"][next_node]["deadline"])
        )
        document.due_date = document.workflow["nodes"][next_node]["due_date"]
    except:
        pass
    db.session.add(document)
    db.session.commit()
    return document, 0


def previous_status_controller(document_id):
    document = get_document_controller(document_id)
    workflow = document.workflow

    current_node = workflow["current_node"]
    if current_node is None:
        return document, 1

    previous_node = None
    nodes = workflow["nodes"].items()
    for key, value in nodes:
        if value["next_node"] == current_node:
            previous_node = key

    # check if there is a previous node
    if previous_node is None:
        return document, 1

    # need to make a copy to track changes to JSON, otherwise the changes are not updated
    document.workflow = copy.deepcopy(document.workflow)
    document.workflow["current_node"] = previous_node
    document.current_step = document.workflow["nodes"][previous_node]["title"]
    try:
        document.workflow["nodes"][previous_node]["due_date"] = str(
            datetime.today()
            + timedelta(days=document.workflow["nodes"][previous_node]["deadline"])
        )
        document.due_date = document.workflow["nodes"][previous_node]["due_date"]
    except:
        pass
    db.session.add(document)
    db.session.commit()
    return document, 0


def document_creation_email_controller(title, company_id):
    company_users = User.query.filter_by(company_id=company_id, active=True)
    email_list = []
    for user in company_users:
        email_list.append(user.email)
    response = send_email_controller(
        "app@lawing.com.br",
        email_list,
        "New Document created",
        title,
        "d-83efa7b8d2fb4742a69dd9059324e148",
    )
    return response


def workflow_status_change_email_controller(document_id, name):
    document = get_document_controller(document_id)
    workflow = document.workflow
    node = workflow["current_node"]
    if node is None:
        return
    users_IDS = workflow["nodes"][node]["responsible_users"]
    status = workflow["nodes"][node]["title"]
    title = document.title
    email_list = []

    # for each user id, add respective user email to email list
    for user_id in users_IDS:
        if isinstance(user_id, int):
            email = User.query.filter_by(id=user_id).first().email
            if email not in email_list:
                email_list.append(email)
    if len(email_list) == 0:
        return
    response = send_email_controller(
        "app@lawing.com.br",
        email_list,
        f"O Documento {title} mudou para o status {status}.",
        name,
        "d-d869f27633274db3810abaa3b60f1833",
    )
    return response


def send_email_controller(
    sender_email, recipient_emails, email_subject, name, template_id
):
    message = Mail(from_email=sender_email, to_emails=recipient_emails)
    # faz as substituições necessárias no template
    if name is not None:
        message.dynamic_template_data = {"name": name}
    message.template_id = template_id
    sg = SendGridAPIClient(current_app.config["SENDGRID_API_KEY"])
    response = sg.send(message)
    return response


def get_download_url_controller(document):
    remote_document = RemoteDocument()
    url = remote_document.download_signed_document(document)
    return url


def get_pdf_download_url_controller(document, version_id):
    remote_document = RemoteDocument()
    url = remote_document.download_pdf_document(document, version_id)
    return url


def get_docx_download_url_controller(document):
    remote_document = RemoteDocument()
    url = remote_document.download_docx_document(document)
    return url


def fill_signing_date_controller(document, text):
    signing_date = json.dumps(date.today().strftime("%d/%m/%Y"), default=str).replace(
        '"', " "
    )
    variable = {"CURRENT_DATE": signing_date}
    document.variables["SIGN_DATE"] = signing_date
    db.session.add(document)
    db.session.commit()

    if document.text_type == ".txt":
        filled_text = fill_text_with_variables(text, variable)
    elif document.text_type == ".docx":
        filled_text = fill_docx_with_variables(document, text, variable)

    return filled_text


def convert_pdf_controller(text):
    remote_document = RemoteDocument()
    template = remote_document.get_template()
    formatted_text = fill_text_with_variables(template, {"text_contract": text})
    document_pdf = requests.post(
        current_app.config["HTMLTOPDF_API_URL"], data=formatted_text.encode("utf-8")
    ).content
    return base64.b64decode(document_pdf)


def change_variables_controller(document, new_variables, email, variables, draft=False):
    document_template = DocumentTemplate.query.filter_by(
        id=document.document_template_id
    ).first()
    current_user = User.query.filter_by(email=email).first()
    current_date = datetime.now().astimezone().replace(microsecond=0).isoformat()
    versions = document.versions
    current_version = int(versions[0]["id"])
    new_version = current_version + 1
    version = {
        "description": "Variáveis do documento foram editadas",
        "email": email,
        "created_at": current_date,
        "id": str(new_version),
        "comments": "",
        "created_by": current_user.name,
    }

    # need to make a copy to track changes to JSON, otherwise the changes are not updated
    document.versions = copy.deepcopy(document.versions)
    document.versions.insert(0, version)
    document.variables = variables

    # Document is being finished, create remote document and update title/draft
    if not draft and document.draft:
        document.draft = draft
        document.title = document.title.replace("Rascunho: ", "")
        try:
            create_remote_document(
                document, variables, document_template, document.company_id
            )
        except:
            raise BadRequest(
                description="Something went wrong while creating the remote document"
            )
    # Document is already finished and is updating it's variables
    elif not document.draft:
        update_variables(
            document, document_template, document.company_id, new_variables
        )
    db.session.add(document)
    db.session.commit()


def edit_document_workflow_controller(
    document,
    new_group: dict = None,
    new_responsible_users: list = None,
    new_due_date: str = None,
):
    new_workflow = copy.deepcopy(document.workflow)
    current_node = document.current_step
    new_node = new_workflow["nodes"][current_node]

    if new_group and new_group.get("id", None):
        new_node["responsible_group"] = format_workflow_responsible_group(new_group)

    if new_responsible_users:
        group = new_group if new_group else new_node["responsible_group"]
        new_node["responsible_users"] = format_workflow_responsible_users(
            new_responsible_users, group
        )

    # Overwrite due_date column with new_workflow due_date for possible changes
    if new_due_date:
        new_due_date_converted = datetime.fromisoformat(new_due_date)
        new_node["due_date"] = new_due_date_converted.isoformat()
        document.due_date = new_due_date_converted

    # Copy new_workflow into document and add to the database
    new_workflow["nodes"][current_node] = new_node
    document.workflow = new_workflow
    db.session.add(document)
    db.session.commit()
    return document


def separate_string(filename):
    doc = docxDocument(filename)
    for p in doc.paragraphs:
        if "{{ INSERT_PARAGRAFO }}" in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if "{{ INSERT_PARAGRAFO }}" in inline[i].text:
                    text = inline[i].text.split("{{ INSERT_PARAGRAFO }}")
                    inline[i].text = text[0]
                    inline[i].add_break()
                    for j in range(len(text)):
                        if j != 0:
                            run = p.add_run(text[j])
                            run.add_break()
    doc.save(filename)


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def convert_docx_to_pdf_and_save(document, filled_docx_io):
    convertapi.api_secret = current_app.config["CONVERTAPI_SECRET_KEY"]

    upload_io = convertapi.UploadIO(filled_docx_io, "filled_docx_io.docx")

    result = convertapi.convert("pdf", {"File": upload_io}, from_format="docx")

    response = requests.get(result.response["Files"][0]["Url"])

    pdf_io = io.BytesIO(response.content)

    remote_path = f'{document.company_id}/{current_app.config["AWS_S3_DOCUMENTS_ROOT"]}/{document.id}/{document.versions[0]["id"]}.pdf'
    remote_document = RemoteDocument()
    remote_document.upload_pdf(pdf_io, remote_path)


def update_variables(document, document_template, company_id, variables):
    remote_document = RemoteDocument()
    docx_io = remote_document.download_docx_from_template(document_template, company_id)
    filled_docx_io = fill_docx_with_variables(document, docx_io, variables)

    convert_docx_to_pdf_and_save(document, filled_docx_io)
    remote_document.upload_filled_docx_to_documents(
        document, filled_docx_io, document_template.text_type
    )


def fill_docx_with_variables(document, docx_io, variables):

    doc = docxDocument(docx_io)

    for key in list(variables):
        if key.startswith("image_") and variables[key]:
            img_bytes = base64.decodebytes(
                variables[key].split("base64,")[1].encode("ascii")
            )
            image = io.BytesIO(img_bytes)
            image_obj = Image.open(io.BytesIO(img_bytes))
            proportion = float(image_obj.size[1] / image_obj.size[0])
            for para in doc.paragraphs:
                if key.split("image_")[1] in para.text:
                    for page in document.form:
                        for field in page["fields"]:
                            if field["variable"]["name"] == key.strip("image_"):
                                width_size = field["variable"].get("width", 8)
                                height_size = width_size * proportion
                                r = para.add_run()
                                r.add_picture(
                                    image, width=Cm(width_size), height=Cm(height_size)
                                )
                                break
    doc.save(docx_io)
    docx_template = DocxTemplate(docx_io)
    docx_template.render(variables)

    filled_text_io = io.BytesIO()
    docx_template.save(filled_text_io)
    filled_text_io.seek(0)

    separate_string(filled_text_io)
    doc = docxDocument(filled_text_io)
    for para in doc.paragraphs:
        if para.text == "" or len(para.text) < 2:
            soup = BeautifulSoup(para._p.xml, "xml")
            if len(soup.find_all("w:numId")) > 0:
                delete_paragraph(para)

    fileobj = io.BytesIO()
    doc.save(fileobj)
    fileobj.seek(0)

    return fileobj


def fill_text_with_variables(text_template, variables):
    jinja_template = jinja_env.from_string(text_template.decode())
    filled_text = jinja_template.render(variables)

    return filled_text


def validate_document_ids_controller(document_ids, company_id):
    valid_ids = db.session.scalars(
        Document.query.filter(Document.id.in_(document_ids))
        .filter_by(company_id=company_id)
        .with_entities(Document.id)
    ).all()
    return valid_ids


def delete_multiple_documents_controller(document_ids):
    Document.query.filter(Document.id.in_(document_ids)).update(
        {"deleted": True, "deleted_at": datetime.utcnow()}
    )
    db.session.commit()


def undelete_document_controller(document):
    document.deleted = False
    document.deleted_at = None
    db.session.commit()


def create_remote_document(document, variables, document_template, company_id):
    company = Company.query.filter_by(id=company_id).first()

    current_date_dict = get_current_date_dict()
    remote_document = RemoteDocument()
    variables = copy.deepcopy(variables)
    variables = format_variables(variables, document_template.id)
    variables.update(current_date_dict)
    if document_template.text_type == ".txt":
        text_template = remote_document.download_text_from_template(document)
        filled_text = fill_text_with_variables(text_template, variables).encode()
        remote_document.upload_filled_text_to_documents(document, filled_text)
    else:
        docx_io = remote_document.download_docx_from_template(
            document_template, company_id
        )
        filled_docx_io = fill_docx_with_variables(document, docx_io, variables)
        convert_docx_to_pdf_and_save(document, filled_docx_io)
        remote_document.upload_filled_docx_to_documents(
            document, filled_docx_io, document_template.text_type
        )

    # Call webhooks listed for the company
    webhook_list = Webhook.query.filter_by(company_id=company_id).all()
    sns = boto3.client("sns")
    topic_arn = current_app.config["SNS_NEWDOCUMENT_ARN"]
    for webhook in webhook_list:
        if webhook.pdf:
            webhook_text_type = ".pdf"
        elif webhook.docx:
            webhook_text_type = ".docx"
        else:
            webhook_text_type = None

        message = {
            "document_title": document.title,
            "document_id": document.id,
            "document_template_name": document_template.name,
            "document_template_id": document_template.id,
            "company_name": company.name,
            "company_id": company.id,
            "user_id": document.user.id,
            "user_name": document.user.name,
            "created_at": json.dumps(
                document.created_at, indent=4, sort_keys=True, default=str
            ),
            "valor_contrato": document.valor_contrato,
            "data_inicio_contrato": document.data_inicio_contrato,
            "data_final_contrato": document.data_final_contrato,
            "data_assinatura": document.data_assinatura,
            "nome_contrato": document.nome_contrato,
            "webhook_url": webhook.webhook,
            "text_type": webhook_text_type,
        }

        response = sns.publish(
            TargetArn=topic_arn,
            Message=json.dumps({"default": json.dumps(message)}),
            MessageStructure="json",
        )
